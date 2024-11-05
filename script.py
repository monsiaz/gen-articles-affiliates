import openai
import os
from io import BytesIO
from docx import Document
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
from google.oauth2 import service_account

# Load OpenAI API key from file
with open('/Users/simonazoulay/PBN_GEN_TEXT_AFF/keys.txt', 'r') as file:
    openai.api_key = file.readline().strip()

# Configuration variables
num_titles = 8
subject = "Bilan de société - Le bilan financier est un document comptable obligatoire dans chaque entreprise. On le décrit comme une image à un instant T du patrimoine et des ressources d’une entreprise. Il permet d’illustrer la situation économique de cette dernière en fonction de son actif et de son passif."
project_name = "bilan-societe.fr"
temperature = 0.4
drive_folder_id = "1M2OhC0woP1fUbixjiGVwowKj3pdR1Oql"

# Create a new folder in Google Drive
def create_drive_folder(service, folder_name, parent_folder_id):
    file_metadata = {
        'name': folder_name,
        'mimeType': 'application/vnd.google-apps.folder',
        'parents': [parent_folder_id]
    }
    folder = service.files().create(body=file_metadata, fields='id').execute()
    return folder.get('id')

# Generate a list of unique titles for the given subject
def generate_article_titles(subject, num_titles, temperature):
    prompt = f"""
    Ton objectif est de générer {num_titles} sujets d'articles uniques et bien distincts relatifs à {subject}.

    Ces articles doivent répondre aux préoccupations des lecteurs professionnels français, avec une valeur ajoutée tangible dans les domaines économiques, juridiques ou organisationnels. Chaque sujet doit être original et offrir un angle différent, sans recoupement avec les autres titres.

    Ecris à la française, en ne mettant la majuscule dans les titres que sur le premier mot, sauf pour les noms propres, les acronymes et tout autre mot qui requiert une ou des majuscules.
    
    Génère {num_titles} titres d'articles variés, granularisés, et sans redondance. Aucun titre ne doit être identique ni trop similaire aux autres. Le sujet doit rester en lien avec le monde de l'entreprise et pas de hors-sujet trop large ou trop éloignés de {subject}.
    """

    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "Tu es un responsable éditorial spécialisé dans le domaine entrepreneurial, légal et financier, et tu génères des sujets d'articles pertinents et granulaires."},
            {"role": "user", "content": prompt}
        ],
        temperature=temperature
    )

    titles = list(set(response.choices[0].message.content.strip().split("\n")))  # Remove duplicates
    if len(titles) < num_titles:
        titles += generate_article_titles(subject, num_titles - len(titles), temperature)  # Generate additional titles if duplicates
    return titles[:num_titles]

# Generate a detailed outline for each title
def generate_detailed_outline(title, global_context):
    outline_prompt = f"""
    Pour le sujet suivant : "{title}", génère un plan très détaillé et logique pour l'article en tenant compte du contexte suivant :
    {global_context}

    Le plan doit inclure les principales sections et sous-sections nécessaires, des idées pour chaque sous-section, et une progression logique des idées. Garde en tête que l'article est destiné à un public professionnel français.
    """

    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "Tu es un expert rédacteur spécialisé dans la création de contenu structuré pour un public professionnel dans les domaines juridique et économique."},
            {"role": "user", "content": outline_prompt}
        ],
        temperature=0.5  # Légère augmentation pour encourager la diversité d'idées
    )

    return response.choices[0].message.content.strip()

# Extract the first 500 words from the generated HTML article
def extract_first_500_words(article_html):
    words = article_html.split()
    return " ".join(words[:500]) if len(words) >= 500 else article_html

# Generate an article in HTML for the given title using the outline
def generate_article_html(title, outline, temperature):
    min_word_count = 2200  # Assure un contenu suffisamment long
    article_html = ""
    prompt = f"""
    Rédige un article d'au moins {min_word_count} mots sur le sujet suivant : "{title}" en suivant le plan détaillé ci-dessous.
    
    Plan détaillé :
    {outline}

    Cet article doit être structuré clairement et inclure des balises HTML selon l'exemple suivant :
    
        <h1>Introduction (sans mentionner explicitement le mot "Introduction" - faire un titre d'introduction fin et élaboré)</h1>
        <p>Contenu introductif en quelques phrases.</p>
        <h2>Sous-titre 1</h2>
        <p>Contenu du sous-titre.</p>
        <h3>Sous-sous-titre</h3>
        <p>Détails et explications.</p>
        <h3>Sous-sous-titre</h3>
        <h3>Sous-sous-titre</h3>
        <h2>Sous-titre 2</h2>
        <p>Contenu du sous-titre.</p>
        <h3>Sous-sous-titre</h3>
        <p>Détails et explications.</p>
        <h3>Sous-sous-titre</h3>
        <h3>Sous-sous-titre</h3>
        etc......

        Utilise des balises <h2> et <h3> en gardant des titres concis, explicites et optimisés pour le SEO. 

        Ecris à la française, en ne mettant la majuscule dans les titres que sur le premier mot, sauf pour les noms propres, les acronymes et tout autre mot qui requiert une ou des majuscules.
        
        Essaie d'avoir une structure équilibrée et un plan pertinent. Essaie d'avoir des formulations naturelles et humaines (te détacher d'un style trop IA).
        
        - Chaque section doit être informative, apportant une réelle valeur ajoutée au lecteur.
        - Les points essentiels doivent être mis en valeur avec des balises <strong>.
        - Utilise des listes à puces si nécessaire.
        - Donne des statistiques et des exemples concrets, avec des explications claires.
        - Le ton doit rester professionnel mais engageant.

        La dernière section, en balise <h2>, doit offrir une perspective future ou une réflexion sans utiliser le mot "conclusion" (pas de conclusion - titre explicitant une conclusion - adoptes une approche plus fine). Et évites les formulations préconçues de type "en conclusion" ou "dans un environnement" dans le dernier paragraphe. Sois créatif et original.

         Utilise les balises HTML appropriées pour le formatage. Tu ne me donnes pas d'en-tête (pas de header) Tu commences directement par le <H1> et pas d'introduction HTML.
    """

    for _ in range(3):  # Limite à 3 itérations pour éviter une boucle infinie
        response = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "Tu es un expert rédacteur dans le monde de l'entreprise, et tu écris des articles de fond riches en informations."},
                {"role": "user", "content": prompt}
            ],
            temperature=temperature
        )
        
        new_content = response.choices[0].message.content.strip()
        article_html += "\n" + new_content  # Ajoute chaque génération de contenu
        word_count = len(article_html.split())
        print(f"Nombre de mots générés pour '{title}': {word_count}")  # Affiche le compteur de mots pour suivi

        # Si la longueur est suffisante, arrêter la boucle
        if word_count >= min_word_count:
            break

        # Prépare un prompt d'amélioration si l'article est encore trop court
        prompt = f"""
        Le contenu suivant est encore trop court et nécessite des informations supplémentaires pour atteindre au moins {min_word_count} mots :
        {article_html}.
        Ajoute des détails et des exemples concrets pour chaque section, en suivant toujours le plan détaillé.
        """
    
    return article_html

# Generate SEO title, description, and URL path with retries if title is missing
def generate_seo_optimizations(title, article_html, retries=3):
    seo_title, seo_description, seo_url = "", "", ""
    
    for attempt in range(retries):
        first_500_words = extract_first_500_words(article_html)
        
        seo_title_prompt = f"""
        En tant qu'expert SEO, génère un titre optimisé pour le SEO de moins de 42 caractères pour l'article suivant : "{title}". Utilise les 500 premiers mots suivants pour optimiser le titre : {first_500_words}.
        """
        
        response_title = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": seo_title_prompt}]
        )
        seo_title = response_title.choices[0].message.content.strip()
        
        if seo_title and "désolé" not in seo_title.lower():
            break
    else:
        raise ValueError("Impossible de générer un titre valide après plusieurs tentatives.")

    seo_description_prompt = f"""
    Génère une description SEO optimisée de moins de 126 caractères pour l'article suivant : "{title}". Utilise les 500 premiers mots pour t'assurer que la description soit pertinente et reflète bien le contenu. Tu me renvoies uniquement la description générée (pas de commentaires IA ou autre).    """
    
    response_description = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": seo_description_prompt}]
    )
    seo_description = response_description.choices[0].message.content.strip()
    
    seo_url_prompt = f"""
    Génère un chemin d'URL optimisé pour le SEO pour l'article suivant : "{title}". Le chemin doit être différent du titre et du H1, tout en restant pertinent pour le sujet. Ne renvoie que le chemin d'URL sous la forme /url-optimisee/.
    """
    
    response_url = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": seo_url_prompt}]
    )
    seo_url = response_url.choices[0].message.content.strip()

    return seo_title, seo_description, seo_url

# Save each article as a Word document in memory and upload it to Google Drive
def upload_article_to_drive(article_html, seo_title, seo_description, seo_url, file_name, service, drive_folder_id):
    doc = Document()
    doc.add_paragraph(f"Titre suggéré : {seo_title}")
    doc.add_paragraph(f"Meta-description suggérée : {seo_description}")
    doc.add_paragraph(f"URL suggérée : {seo_url}")
    doc.add_paragraph("Contenu HTML brut :")
    doc.add_paragraph(article_html)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    file_metadata = {
        'name': f"{file_name}.docx",
        'parents': [drive_folder_id]
    }
    media = MediaIoBaseUpload(file_stream, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
    print(f"File ID: {file.get('id')}")

# Main workflow
def main():
    global_context = "Les articles doivent éviter de répéter les sujets suivants :"
    titles = generate_article_titles(subject, num_titles, temperature)
    SCOPES = ['https://www.googleapis.com/auth/drive.file']
    SERVICE_ACCOUNT_FILE = '/Users/simonazoulay/PBN_GEN_TEXT_AFF/leafy-brace-242115-e657a84d991b.json'

    credentials = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE, scopes=SCOPES)
    service = build('drive', 'v3', credentials=credentials)

    project_folder_id = create_drive_folder(service, project_name, drive_folder_id)
    
    for idx, title in enumerate(titles):
        global_context += f' {title};'
        print(f"Generating outline for title {idx + 1}: {title}")
        
        outline = generate_detailed_outline(title, global_context)
        
        print(f"Generating article for title {idx + 1}: {title}")
        article_html = generate_article_html(title, outline, temperature)
        
        seo_title, seo_description, seo_url = generate_seo_optimizations(title, article_html)
        
        file_name = seo_title.replace(' ', '_')
        upload_article_to_drive(article_html, seo_title, seo_description, seo_url, file_name, service, project_folder_id)

if __name__ == "__main__": 
    main()
